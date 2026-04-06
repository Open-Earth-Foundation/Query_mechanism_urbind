"""Helpers for exporting generated markdown documents to `.docx`."""

from __future__ import annotations

from io import BytesIO
import re

from docx import Document
from docx.document import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph

DOCX_MIME_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
HEADING_PATTERN = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
ORDERED_LIST_PATTERN = re.compile(r"^(\s*)\d+[.)]\s+(.+?)\s*$")
UNORDERED_LIST_PATTERN = re.compile(r"^(\s*)[-*+]\s+(.+?)\s*$")
INLINE_TOKEN_PATTERN = re.compile(
    r"(\*\*[^*]+\*\*|__[^_]+__|`[^`]+`|\*[^*]+\*|_[^_]+_|\[[^\]]+\]\([^)]+\))"
)
TABLE_ALIGNMENT_PATTERN = re.compile(r"^:?-{3,}:?$")
EXPORT_CITATION_PATTERN = re.compile(r"(?:\s*\[ref_[^\]\s]+\](?!\())+")


def markdown_to_docx_bytes(markdown: str) -> bytes:
    """Render markdown text into a `.docx` byte payload."""
    document = Document()
    _configure_document(document)

    lines = markdown.splitlines()
    index = 0
    while index < len(lines):
        raw_line = lines[index]
        stripped = raw_line.strip()
        if not stripped:
            index += 1
            continue
        if stripped.startswith("```"):
            index = _add_fenced_code_block(document, lines, index)
            continue
        if _starts_markdown_table(lines, index):
            index = _add_markdown_table(document, lines, index)
            continue
        heading_match = HEADING_PATTERN.match(raw_line)
        if heading_match:
            _add_heading(document, heading_match.group(2), len(heading_match.group(1)))
            index += 1
            continue
        if _is_list_item(raw_line):
            index = _add_list_block(document, lines, index)
            continue
        if stripped.startswith(">"):
            index = _add_blockquote(document, lines, index)
            continue
        if _is_horizontal_rule(stripped):
            document.add_paragraph()
            index += 1
            continue
        index = _add_paragraph_block(document, lines, index)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _configure_document(document: DocxDocument) -> None:
    """Apply basic spacing defaults for exported reports."""
    document.core_properties.title = "Exported Query Mechanism Document"
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Aptos"
    normal_style.font.size = Pt(11)


def _starts_markdown_table(lines: list[str], index: int) -> bool:
    """Return True when the current line starts a GFM-style table."""
    if index + 1 >= len(lines):
        return False
    header_row = lines[index].strip()
    separator_row = lines[index + 1].strip()
    if "|" not in header_row:
        return False
    separator_cells = _split_table_row(separator_row)
    if not separator_cells:
        return False
    return all(
        TABLE_ALIGNMENT_PATTERN.fullmatch(cell.replace(" ", "")) is not None
        for cell in separator_cells
    )


def _add_markdown_table(
    document: DocxDocument, lines: list[str], start_index: int
) -> int:
    """Append one markdown table and return the next unread line index."""
    header_cells = _split_table_row(lines[start_index])
    rows = [header_cells]
    index = start_index + 2
    while index < len(lines):
        candidate = lines[index].strip()
        if not candidate or "|" not in candidate:
            break
        rows.append(_split_table_row(candidate))
        index += 1

    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"

    for row_index, row in enumerate(rows):
        for column_index in range(column_count):
            cell = table.rows[row_index].cells[column_index]
            cell_text = row[column_index] if column_index < len(row) else ""
            paragraph = cell.paragraphs[0]
            _append_inline_runs(paragraph, cell_text)
            if row_index == 0:
                for run in paragraph.runs:
                    run.bold = True
    return index


def _add_heading(document: DocxDocument, text: str, level: int) -> None:
    """Append a heading paragraph with inline formatting."""
    paragraph = document.add_paragraph(style=f"Heading {max(1, min(level, 6))}")
    _append_inline_runs(paragraph, text.strip())


def _add_list_block(document: DocxDocument, lines: list[str], start_index: int) -> int:
    """Append contiguous list items and return the next unread line index."""
    index = start_index
    while index < len(lines):
        line = lines[index]
        match = ORDERED_LIST_PATTERN.match(line) or UNORDERED_LIST_PATTERN.match(line)
        if match is None:
            break
        indent_level = _indent_level(match.group(1))
        item_lines = [match.group(2).strip()]
        index += 1
        while index < len(lines):
            continuation = lines[index]
            continuation_stripped = continuation.strip()
            if not continuation_stripped:
                break
            if _is_block_boundary(lines, index) or _is_list_item(continuation):
                break
            item_lines.append(continuation_stripped)
            index += 1

        style_name = (
            "List Number"
            if ORDERED_LIST_PATTERN.match(line) is not None
            else "List Bullet"
        )
        paragraph = document.add_paragraph(style=style_name)
        if indent_level > 0:
            paragraph.paragraph_format.left_indent = Inches(0.25 * indent_level)
        _append_inline_runs(paragraph, " ".join(item_lines))
        if index < len(lines) and not lines[index].strip():
            index += 1
    return index


def _add_blockquote(document: DocxDocument, lines: list[str], start_index: int) -> int:
    """Append a contiguous markdown blockquote and return the next unread index."""
    quote_lines: list[str] = []
    index = start_index
    while index < len(lines):
        stripped = lines[index].strip()
        if not stripped.startswith(">"):
            break
        quote_lines.append(stripped[1:].strip())
        index += 1
    paragraph = document.add_paragraph(style="Intense Quote")
    _append_inline_runs(paragraph, " ".join(line for line in quote_lines if line))
    return index


def _add_paragraph_block(
    document: DocxDocument, lines: list[str], start_index: int
) -> int:
    """Append a normal paragraph block and return the next unread line index."""
    paragraph_lines: list[str] = []
    index = start_index
    while index < len(lines):
        stripped = lines[index].strip()
        if not stripped or _is_block_boundary(lines, index):
            break
        paragraph_lines.append(stripped)
        index += 1
    paragraph = document.add_paragraph()
    _append_inline_runs(paragraph, " ".join(paragraph_lines))
    return index


def _add_fenced_code_block(
    document: DocxDocument, lines: list[str], start_index: int
) -> int:
    """Append a fenced code block and return the next unread line index."""
    index = start_index + 1
    code_lines: list[str] = []
    while index < len(lines):
        if lines[index].strip().startswith("```"):
            index += 1
            break
        code_lines.append(lines[index])
        index += 1
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run("\n".join(code_lines))
    run.font.name = "Courier New"
    run.font.size = Pt(10)
    return index


def _append_inline_runs(paragraph: Paragraph, text: str) -> None:
    """Render simple inline markdown formatting into docx runs."""
    text = _strip_export_citations(text)
    cursor = 0
    for match in INLINE_TOKEN_PATTERN.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor : match.start()])
        token = match.group(0)
        run = paragraph.add_run(_normalize_inline_token_text(token))
        if token.startswith(("**", "__")):
            run.bold = True
        elif token.startswith(("*", "_")):
            run.italic = True
        elif token.startswith("`"):
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        elif token.startswith("["):
            run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
            run.underline = True
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def _strip_export_citations(text: str) -> str:
    """Remove inline ``[ref_*]`` citation markers from export text."""
    return EXPORT_CITATION_PATTERN.sub("", text).strip()


def _normalize_inline_token_text(token: str) -> str:
    """Convert one inline markdown token into readable docx text."""
    if token.startswith(("**", "__")):
        return token[2:-2]
    if token.startswith(("`", "*", "_")):
        return token[1:-1]
    if token.startswith("["):
        label, url = token[1:].split("](", maxsplit=1)
        return f"{label} ({url[:-1]})"
    return token


def _split_table_row(row: str) -> list[str]:
    """Split one markdown table row into trimmed cell strings."""
    trimmed = row.strip()
    if trimmed.startswith("|"):
        trimmed = trimmed[1:]
    if trimmed.endswith("|"):
        trimmed = trimmed[:-1]
    if not trimmed:
        return []
    return [cell.strip() for cell in re.split(r"(?<!\\)\|", trimmed)]


def _is_list_item(line: str) -> bool:
    """Return True when a line starts an ordered or unordered list item."""
    return (
        ORDERED_LIST_PATTERN.match(line) is not None
        or UNORDERED_LIST_PATTERN.match(line) is not None
    )


def _indent_level(indent: str) -> int:
    """Approximate list nesting level from leading whitespace."""
    return max(len(indent.replace("\t", "    ")) // 2, 0)


def _is_block_boundary(lines: list[str], index: int) -> bool:
    """Return True when the line at index starts a non-paragraph block."""
    stripped = lines[index].strip()
    if not stripped:
        return True
    return any(
        (
            stripped.startswith("```"),
            stripped.startswith(">"),
            _is_horizontal_rule(stripped),
            _is_list_item(lines[index]),
            HEADING_PATTERN.match(lines[index]) is not None,
            _starts_markdown_table(lines, index),
        )
    )


def _is_horizontal_rule(stripped: str) -> bool:
    """Return True when a line is a markdown horizontal rule."""
    normalized = stripped.replace(" ", "")
    return normalized in {"---", "***", "___"}


__all__ = ["DOCX_MIME_TYPE", "markdown_to_docx_bytes"]
