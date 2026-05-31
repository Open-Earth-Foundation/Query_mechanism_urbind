import json
import os
import shutil
import threading
import time
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient

from backend.api.main import create_app
from backend.api.services.run_store import TERMINAL_STATUSES
from backend.modules.markdown_researcher.services import build_markdown_chunks_for_file
from backend.utils.config import AppConfig
from backend.utils.paths import RunPaths, create_run_paths
from tests.support import build_test_app_config


def _build_config(runs_dir: Path, markdown_dir: Path) -> AppConfig:
    """Build the API run test config with the required agent sections."""
    return build_test_app_config(
        runs_dir=runs_dir,
        markdown_dir=markdown_dir,
    )


def _write_success_artifacts(question: str, run_id: str, config: AppConfig) -> RunPaths:
    paths = create_run_paths(config.runs_dir, run_id, config.orchestrator.context_bundle_name)
    paths.base_dir.mkdir(parents=True, exist_ok=True)

    context_bundle = {
        "markdown": {"status": "success", "excerpts": []},
        "drafts": [],
        "final": str(paths.final_output),
    }
    paths.context_bundle.write_text(
        json.dumps(context_bundle, ensure_ascii=True, indent=2),
        encoding="utf-8",
    )
    rendered_output = f"# Question\n{question}\n\n# Answer\nStub answer"
    paths.final_output.write_text(rendered_output, encoding="utf-8")

    run_log = {
        "run_id": run_id,
        "question": question,
        "status": "completed",
        "started_at": datetime.now(timezone.utc).isoformat(),
        "completed_at": datetime.now(timezone.utc).isoformat(),
        "finish_reason": "completed (write)",
        "artifacts": {
            "context_bundle": str(paths.context_bundle),
            "final_output": str(paths.final_output),
        },
    }
    paths.run_log.write_text(
        json.dumps(run_log, ensure_ascii=True, indent=2), encoding="utf-8"
    )
    return paths


def _write_run_listing_artifacts(
    *,
    question: str,
    run_id: str,
    status: str,
    config: AppConfig,
    finish_reason: str | None = None,
    error: dict[str, str] | None = None,
) -> RunPaths:
    """Write minimal run artifacts for list and diagnostics route tests."""
    paths = create_run_paths(config.runs_dir, run_id, config.orchestrator.context_bundle_name)
    paths.base_dir.mkdir(parents=True, exist_ok=True)
    run_log = {
        "run_id": run_id,
        "question": question,
        "status": status,
        "started_at": datetime.now(timezone.utc).isoformat(),
        "completed_at": datetime.now(timezone.utc).isoformat(),
        "finish_reason": finish_reason,
        "error": error,
        "artifacts": {},
    }
    paths.run_log.write_text(
        json.dumps(run_log, ensure_ascii=True, indent=2),
        encoding="utf-8",
    )
    return paths


def _write_run_listing_artifact(
    runs_dir: Path,
    *,
    run_id: str,
    started_at: datetime,
    question: str | None,
    inputs: dict[str, object] | None = None,
) -> None:
    """Persist the minimal run.json payload needed by run-list endpoint tests."""
    run_dir = runs_dir / run_id
    run_dir.mkdir(parents=True, exist_ok=True)
    payload: dict[str, object] = {
        "run_id": run_id,
        "status": "completed",
        "started_at": started_at.isoformat(),
        "completed_at": started_at.isoformat(),
    }
    if question is not None:
        payload["question"] = question
    if inputs is not None:
        payload["inputs"] = inputs
    (run_dir / "run.json").write_text(
        json.dumps(payload, ensure_ascii=True, indent=2),
        encoding="utf-8",
    )


def _write_config_file(path: Path, config: AppConfig) -> None:
    """Persist one test config as JSON-compatible YAML."""
    path.write_text(
        json.dumps(config.model_dump(mode="json"), ensure_ascii=True, indent=2),
        encoding="utf-8",
    )


def _write_started_artifacts_with_error_log_input(
    question: str,
    run_id: str,
    config: AppConfig,
) -> RunPaths:
    paths = create_run_paths(config.runs_dir, run_id, config.orchestrator.context_bundle_name)
    paths.base_dir.mkdir(parents=True, exist_ok=True)
    run_log_payload = {
        "run_id": run_id,
        "question": question,
        "status": "started",
        "started_at": datetime.now(timezone.utc).isoformat(),
        "completed_at": None,
        "decisions": [],
        "artifacts": {},
    }
    paths.run_log.write_text(
        json.dumps(run_log_payload, ensure_ascii=True, indent=2),
        encoding="utf-8",
    )
    (paths.base_dir / "run.log").write_text(
        "\n".join(
            [
                "2026-01-01 00:00:00 worker.py:10 - INFO - setup",
                "2026-01-01 00:00:01 worker.py:11 - ERROR - writer crashed",
                "2026-01-01 00:00:02 worker.py:12 - CRITICAL - aborting run",
            ]
        ),
        encoding="utf-8",
    )
    return paths


def _write_failed_artifacts_with_decision_error(
    *,
    question: str,
    run_id: str,
    config: AppConfig,
    finish_reason: str,
    error_code: str,
    error_message: str,
) -> RunPaths:
    """Persist a failed run.json with a structured decision error."""
    paths = create_run_paths(config.runs_dir, run_id, config.orchestrator.context_bundle_name)
    paths.base_dir.mkdir(parents=True, exist_ok=True)
    run_log_payload = {
        "run_id": run_id,
        "question": question,
        "status": "failed",
        "started_at": datetime.now(timezone.utc).isoformat(),
        "completed_at": datetime.now(timezone.utc).isoformat(),
        "finish_reason": finish_reason,
        "decisions": [
            {
                "status": "error",
                "reason": "Persisted pipeline failure",
                "error": {
                    "code": error_code,
                    "message": error_message,
                },
            }
        ],
        "artifacts": {},
    }
    paths.run_log.write_text(
        json.dumps(run_log_payload, ensure_ascii=True, indent=2),
        encoding="utf-8",
    )
    (paths.base_dir / "run.log").write_text(
        "2026-01-01 00:00:01 worker.py:11 - ERROR - persisted pipeline failure",
        encoding="utf-8",
    )
    return paths


def _write_markdown_reference_artifacts(
    paths: RunPaths,
    references_payload: dict[str, object] | None = None,
    excerpts_payload: dict[str, object] | None = None,
) -> None:
    paths.markdown_dir.mkdir(parents=True, exist_ok=True)
    if references_payload is not None:
        paths.markdown_references.write_text(
            json.dumps(references_payload, ensure_ascii=True, indent=2),
            encoding="utf-8",
        )
    if excerpts_payload is not None:
        paths.markdown_excerpts.write_text(
            json.dumps(excerpts_payload, ensure_ascii=True, indent=2),
            encoding="utf-8",
        )


def _write_markdown_batches_artifact(paths: RunPaths, payload: dict[str, object]) -> None:
    """Persist markdown batch metadata for chunk-to-path lookup tests."""
    paths.markdown_dir.mkdir(parents=True, exist_ok=True)
    (paths.markdown_dir / "batches.json").write_text(
        json.dumps(payload, ensure_ascii=True, indent=2),
        encoding="utf-8",
    )


def _poll_until_terminal(
    client: TestClient,
    run_id: str,
    timeout_seconds: float = 3.0,
) -> dict[str, object]:
    deadline = time.monotonic() + timeout_seconds
    last_payload: dict[str, object] = {}
    while time.monotonic() < deadline:
        response = client.get(f"/api/v1/runs/{run_id}/status")
        assert response.status_code == 200
        payload = response.json()
        last_payload = payload
        if payload["status"] in TERMINAL_STATUSES:
            return payload
        time.sleep(0.02)
    raise AssertionError(f"Run {run_id} did not reach terminal status: {last_payload}")


def test_api_run_lifecycle_success(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    runs_dir = tmp_path / "output"
    markdown_dir = tmp_path / "documents"
    markdown_dir.mkdir(parents=True, exist_ok=True)

    def _stub_load_config(_path: Path | None = None) -> AppConfig:
        return _build_config(runs_dir=runs_dir, markdown_dir=markdown_dir)

    def _stub_run_pipeline(
        question: str,
        config: AppConfig,
        run_id: str | None = None,
        log_llm_payload: bool = True,
        analysis_mode: str = "aggregate",
        api_key_override: str | None = None,
        selected_cities: list[str] | None = None,
    ) -> RunPaths:
        assert run_id is not None
        assert isinstance(log_llm_payload, bool)
        assert analysis_mode == "aggregate"
        assert api_key_override is None
        assert selected_cities is None
        return _write_success_artifacts(question=question, run_id=run_id, config=config)

    monkeypatch.setattr("backend.api.services.run_executor.load_config", _stub_load_config)
    monkeypatch.setattr("backend.api.services.run_executor.run_pipeline", _stub_run_pipeline)

    app = create_app(runs_dir=runs_dir, max_workers=2)
    with TestClient(app) as client:
        start = client.post(
            "/api/v1/runs",
            json={"question": "What are the key initiatives?", "run_id": "run-success"},
        )
        assert start.status_code == 202
        start_payload = start.json()
        assert start_payload["run_id"] == "run-success"

        terminal = _poll_until_terminal(client, "run-success")
        assert terminal["status"] == "completed"

        output_response = client.get("/api/v1/runs/run-success/output")
        assert output_response.status_code == 200
        output_payload = output_response.json()
        assert output_payload["status"] == "completed"
        assert "Stub answer" in output_payload["content"]

        context_response = client.get("/api/v1/runs/run-success/context")
        assert context_response.status_code == 200
        context_payload = context_response.json()
        assert context_payload["status"] == "completed"
        assert isinstance(context_payload["context_bundle"], dict)


def test_api_run_lifecycle_dev_mode_ignores_blank_optional_queries(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """API runs should trim optional direct queries and omit blank ones."""
    runs_dir = tmp_path / "output"
    markdown_dir = tmp_path / "documents"
    markdown_dir.mkdir(parents=True, exist_ok=True)

    def _stub_load_config(_path: Path | None = None) -> AppConfig:
        return _build_config(runs_dir=runs_dir, markdown_dir=markdown_dir)

    def _stub_run_pipeline(
        question: str,
        config: AppConfig,
        run_id: str | None = None,
        log_llm_payload: bool = True,
        analysis_mode: str = "aggregate",
        query_mode: str = "standard",
        query_2: str | None = None,
        query_3: str | None = None,
        api_key_override: str | None = None,
        selected_cities: list[str] | None = None,
    ) -> RunPaths:
        assert run_id is not None
        assert isinstance(log_llm_payload, bool)
        assert analysis_mode == "aggregate"
        assert query_mode == "dev"
        assert query_2 == "retrofit milestones and deadlines"
        assert query_3 is None
        assert api_key_override is None
        assert selected_cities is None
        return _write_success_artifacts(question=question, run_id=run_id, config=config)

    monkeypatch.setattr("backend.api.services.run_executor.load_config", _stub_load_config)
    monkeypatch.setattr("backend.api.services.run_executor.run_pipeline", _stub_run_pipeline)

    app = create_app(runs_dir=runs_dir, max_workers=2)
    with TestClient(app) as client:
        start = client.post(
            "/api/v1/runs",
            json={
                "question": "What are the key retrofit initiatives?",
                "run_id": "run-dev-mode",
                "query_mode": "dev",
                "query_2": "  retrofit milestones and deadlines  ",
                "query_3": "   ",
            },
        )
        assert start.status_code == 202
        terminal = _poll_until_terminal(client, "run-dev-mode")
        assert terminal["status"] == "completed"


def test_api_run_lifecycle_standard_mode_passes_optional_queries(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """Standard API runs should pass only user-provided optional queries."""
    runs_dir = tmp_path / "output"
    markdown_dir = tmp_path / "documents"
    markdown_dir.mkdir(parents=True, exist_ok=True)

    def _stub_load_config(_path: Path | None = None) -> AppConfig:
        return _build_config(runs_dir=runs_dir, markdown_dir=markdown_dir)

    def _stub_run_pipeline(
        question: str,
        config: AppConfig,
        run_id: str | None = None,
        log_llm_payload: bool = True,
        analysis_mode: str = "aggregate",
        query_2: str | None = None,
        query_3: str | None = None,
        api_key_override: str | None = None,
        selected_cities: list[str] | None = None,
    ) -> RunPaths:
        assert question == "Compare public EV charging targets exactly as entered."
        assert run_id is not None
        assert isinstance(log_llm_payload, bool)
        assert analysis_mode == "aggregate"
        assert query_2 == "charging rollout milestones"
        assert query_3 is None
        assert api_key_override is None
        assert selected_cities is None
        return _write_success_artifacts(question=question, run_id=run_id, config=config)

    monkeypatch.setattr("backend.api.services.run_executor.load_config", _stub_load_config)
    monkeypatch.setattr("backend.api.services.run_executor.run_pipeline", _stub_run_pipeline)

    app = create_app(runs_dir=runs_dir, max_workers=2)
    with TestClient(app) as client:
        start = client.post(
            "/api/v1/runs",
            json={
                "question": "Compare public EV charging targets exactly as entered.",
                "run_id": "run-standard-optional-query",
                "query_2": "  charging rollout milestones  ",
                "query_3": "",
            },
        )
        assert start.status_code == 202
        terminal = _poll_until_terminal(client, "run-standard-optional-query")
        assert terminal["status"] == "completed"


def test_api_get_run_reference_returns_record_from_references_artifact(
    tmp_path: Path,
) -> None:
    runs_dir = tmp_path / "output"
    markdown_dir = tmp_path / "documents"
    markdown_dir.mkdir(parents=True, exist_ok=True)
    config = _build_config(runs_dir=runs_dir, markdown_dir=markdown_dir)
    paths = _write_success_artifacts(
        question="Reference lookup run",
        run_id="run-reference",
        config=config,
    )
    _write_markdown_reference_artifacts(
        paths=paths,
        references_payload={
            "run_id": "run-reference",
            "reference_count": 1,
            "references": [
                {
                    "ref_id": "ref_1",
                    "excerpt_index": 0,
                    "city_name": "Leipzig",
                    "quote": "Leipzig plans to expand charging infrastructure.",
                    "partial_answer": "Leipzig plans charging infrastructure expansion.",
                    "source_chunk_ids": ["chunk_leipzig_1"],
                }
            ],
        },
    )

    app = create_app(runs_dir=runs_dir, max_workers=1, markdown_dir=markdown_dir)
    with TestClient(app) as client:
        response = client.get("/api/v1/runs/run-reference/references/ref_1")
        assert response.status_code == 200
        payload = response.json()
        assert payload["run_id"] == "run-reference"
        assert payload["ref_id"] == "ref_1"
        assert payload["excerpt_index"] == 0
        assert payload["city_name"] == "Leipzig"
        assert payload["source_chunk_ids"] == ["chunk_leipzig_1"]


def test_api_list_run_references_returns_lightweight_payload_by_default(
    tmp_path: Path,
) -> None:
    runs_dir = tmp_path / "output"
    markdown_dir = tmp_path / "documents"
    markdown_dir.mkdir(parents=True, exist_ok=True)
    config = _build_config(runs_dir=runs_dir, markdown_dir=markdown_dir)
    paths = _write_success_artifacts(
        question="Reference list run",
        run_id="run-reference-list",
        config=config,
    )
    _write_markdown_reference_artifacts(
        paths=paths,
        references_payload={
            "run_id": "run-reference-list",
            "reference_count": 1,
            "references": [
                {
                    "ref_id": "ref_1",
                    "excerpt_index": 0,
                    "city_name": "Leipzig",
                    "quote": "Leipzig plans to expand charging infrastructure.",
                    "partial_answer": "Leipzig plans charging infrastructure expansion.",
                    "source_chunk_ids": ["chunk_leipzig_1"],
                }
            ],
        },
    )

    app = create_app(runs_dir=runs_dir, max_workers=1, markdown_dir=markdown_dir)
    with TestClient(app) as client:
        response = client.get("/api/v1/runs/run-reference-list/references")
        assert response.status_code == 200
        payload = response.json()
        assert payload["run_id"] == "run-reference-list"
        assert payload["reference_count"] == 1
        item = payload["references"][0]
        assert item["ref_id"] == "ref_1"
        assert item["city_name"] == "Leipzig"
        assert "quote" not in item
        assert "source_chunk_ids" not in item


def test_api_list_run_references_supports_ref_filter_and_include_quote(
    tmp_path: Path,
) -> None:
    runs_dir = tmp_path / "output"
    markdown_dir = tmp_path / "documents"
    markdown_dir.mkdir(parents=True, exist_ok=True)
    config = _build_config(runs_dir=runs_dir, markdown_dir=markdown_dir)
    paths = _write_success_artifacts(
        question="Reference list filter run",
        run_id="run-reference-list-filter",
        config=config,
    )
    _write_markdown_reference_artifacts(
        paths=paths,
        references_payload={
            "run_id": "run-reference-list-filter",
            "reference_count": 2,
            "references": [
                {
                    "ref_id": "ref_1",
                    "excerpt_index": 0,
                    "city_name": "Leipzig",
                    "quote": "Leipzig plans charging.",
                    "partial_answer": "Leipzig plans charging.",
                    "source_chunk_ids": ["chunk_leipzig_1"],
                },
                {
                    "ref_id": "ref_2",
                    "excerpt_index": 1,
                    "city_name": "Munich",
                    "quote": "Munich reports chargers.",
                    "partial_answer": "Munich reports chargers.",
                    "source_chunk_ids": ["chunk_munich_1"],
                },
            ],
        },
    )

    app = create_app(runs_dir=runs_dir, max_workers=1, markdown_dir=markdown_dir)
    with TestClient(app) as client:
        response = client.get(
            "/api/v1/runs/run-reference-list-filter/references?ref_id=ref_2&include_quote=true"
        )
        assert response.status_code == 200
        payload = response.json()
        assert payload["reference_count"] == 1
        item = payload["references"][0]
        assert item["ref_id"] == "ref_2"
        assert item["quote"] == "Munich reports chargers."
        assert item["source_chunk_ids"] == ["chunk_munich_1"]


def test_api_get_run_reference_returns_not_found_for_unknown_ref(
    tmp_path: Path,
) -> None:
    runs_dir = tmp_path / "output"
    markdown_dir = tmp_path / "documents"
    markdown_dir.mkdir(parents=True, exist_ok=True)
    config = _build_config(runs_dir=runs_dir, markdown_dir=markdown_dir)
    paths = _write_success_artifacts(
        question="Unknown reference lookup run",
        run_id="run-reference-missing",
        config=config,
    )
    _write_markdown_reference_artifacts(
        paths=paths,
        references_payload={
            "run_id": "run-reference-missing",
            "reference_count": 1,
            "references": [
                {
                    "ref_id": "ref_1",
                    "excerpt_index": 0,
                    "city_name": "Leipzig",
                    "quote": "Leipzig plans to expand charging infrastructure.",
                    "partial_answer": "Leipzig plans charging infrastructure expansion.",
                    "source_chunk_ids": ["chunk_leipzig_1"],
                }
            ],
        },
    )

    app = create_app(runs_dir=runs_dir, max_workers=1, markdown_dir=markdown_dir)
    with TestClient(app) as client:
        response = client.get("/api/v1/runs/run-reference-missing/references/ref_2")
        assert response.status_code == 404


def test_api_get_run_reference_falls_back_to_excerpts_when_references_missing(
    tmp_path: Path,
) -> None:
    runs_dir = tmp_path / "output"
    markdown_dir = tmp_path / "documents"
    markdown_dir.mkdir(parents=True, exist_ok=True)
    config = _build_config(runs_dir=runs_dir, markdown_dir=markdown_dir)
    paths = _write_success_artifacts(
        question="Fallback reference lookup run",
        run_id="run-reference-fallback",
        config=config,
    )
    _write_markdown_reference_artifacts(
        paths=paths,
        excerpts_payload={
            "status": "success",
            "excerpts": [
                {
                    "city_name": "Munich",
                    "quote": "Munich has 43 charging points as of 2024.",
                    "partial_answer": "Munich reports 43 charging points as of 2024.",
                    "source_chunk_ids": ["chunk_munich_1"],
                }
            ],
        },
    )

    app = create_app(runs_dir=runs_dir, max_workers=1, markdown_dir=markdown_dir)
    with TestClient(app) as client:
        response = client.get("/api/v1/runs/run-reference-fallback/references/ref_1")
        assert response.status_code == 200
        payload = response.json()
        assert payload["ref_id"] == "ref_1"
        assert payload["city_name"] == "Munich"
        assert payload["source_chunk_ids"] == ["chunk_munich_1"]

        list_response = client.get("/api/v1/runs/run-reference-fallback/references")
        assert list_response.status_code == 200
        list_payload = list_response.json()
        assert list_payload["reference_count"] == 1
        assert list_payload["references"][0]["ref_id"] == "ref_1"


def test_api_get_run_reference_rejects_invalid_ref_format(tmp_path: Path) -> None:
    runs_dir = tmp_path / "output"
    markdown_dir = tmp_path / "documents"
    markdown_dir.mkdir(parents=True, exist_ok=True)
    config = _build_config(runs_dir=runs_dir, markdown_dir=markdown_dir)
    _write_success_artifacts(
        question="Invalid reference id run",
        run_id="run-reference-invalid",
        config=config,
    )

    app = create_app(runs_dir=runs_dir, max_workers=1, markdown_dir=markdown_dir)
    with TestClient(app) as client:
        response = client.get("/api/v1/runs/run-reference-invalid/references/ref_x")
        assert response.status_code == 400
        query_response = client.get(
            "/api/v1/runs/run-reference-invalid/references?ref_id=ref_x"
        )
        assert query_response.status_code == 400


def test_api_list_run_source_chunks_returns_full_chunk_content(tmp_path: Path) -> None:
    runs_dir = tmp_path / "output"
    markdown_dir = tmp_path / "documents"
    markdown_dir.mkdir(parents=True, exist_ok=True)
    config = _build_config(runs_dir=runs_dir, markdown_dir=markdown_dir)
    config_path = tmp_path / "llm_config.yaml"
    _write_config_file(config_path, config)

    source_path = markdown_dir / "Leipzig.md"
    source_content = (
        "# Leipzig\n\n"
        "Leipzig plans to expand charging infrastructure across municipal districts.\n"
    )
    source_path.write_text(source_content, encoding="utf-8")
    chunks = build_markdown_chunks_for_file(source_path, config.markdown_researcher)
    assert len(chunks) == 1
    chunk_id = str(chunks[0]["chunk_id"])

    paths = _write_success_artifacts(
        question="Source chunk lookup run",
        run_id="run-source-chunks",
        config=config,
    )
    _write_markdown_batches_artifact(
        paths=paths,
        payload={
            "batches": [
                {
                    "city_name": "leipzig",
                    "batch_index": 1,
                    "chunk_count": 1,
                    "estimated_tokens": 10,
                    "chunks": [
                        {
                            "chunk_id": chunk_id,
                            "path": str(source_path),
                            "chunk_index": 0,
                        }
                    ],
                }
            ]
        },
    )

    app = create_app(
        runs_dir=runs_dir,
        max_workers=1,
        markdown_dir=markdown_dir,
        config_path=config_path,
    )
    with TestClient(app) as client:
        response = client.get(
            f"/api/v1/runs/run-source-chunks/source-chunks?chunk_id={chunk_id}"
        )
        assert response.status_code == 200
        payload = response.json()
        assert payload["run_id"] == "run-source-chunks"
        assert payload["chunk_count"] == 1
        assert payload["chunks"][0]["chunk_id"] == chunk_id
        assert payload["chunks"][0]["content"] == source_content


def test_api_list_run_source_chunks_returns_not_found_for_missing_chunk(
    tmp_path: Path,
) -> None:
    runs_dir = tmp_path / "output"
    markdown_dir = tmp_path / "documents"
    markdown_dir.mkdir(parents=True, exist_ok=True)
    config = _build_config(runs_dir=runs_dir, markdown_dir=markdown_dir)
    config_path = tmp_path / "llm_config.yaml"
    _write_config_file(config_path, config)
    _write_success_artifacts(
        question="Missing chunk lookup run",
        run_id="run-source-chunks-missing",
        config=config,
    )

    app = create_app(
        runs_dir=runs_dir,
        max_workers=1,
        markdown_dir=markdown_dir,
        config_path=config_path,
    )
    with TestClient(app) as client:
        response = client.get(
            "/api/v1/runs/run-source-chunks-missing/source-chunks?chunk_id=chunk_missing"
        )
        assert response.status_code == 404


def test_api_list_run_source_chunks_reuses_cached_config_until_file_changes(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """Source chunk requests reuse cached config and reload after config mtime changes."""
    runs_dir = tmp_path / "output"
    markdown_dir = tmp_path / "documents"
    markdown_dir.mkdir(parents=True, exist_ok=True)
    config = _build_config(runs_dir=runs_dir, markdown_dir=markdown_dir)
    config_path = tmp_path / "llm_config.yaml"
    _write_config_file(config_path, config)

    source_path = markdown_dir / "Leipzig.md"
    source_content = (
        "# Leipzig\n\n"
        "Leipzig plans to expand charging infrastructure across municipal districts.\n"
    )
    source_path.write_text(source_content, encoding="utf-8")
    chunks = build_markdown_chunks_for_file(source_path, config.markdown_researcher)
    chunk_id = str(chunks[0]["chunk_id"])

    paths = _write_success_artifacts(
        question="Source chunk cache lookup run",
        run_id="run-source-chunks-cache",
        config=config,
    )
    _write_markdown_batches_artifact(
        paths=paths,
        payload={
            "batches": [
                {
                    "city_name": "leipzig",
                    "batch_index": 1,
                    "chunk_count": 1,
                    "estimated_tokens": 10,
                    "chunks": [
                        {
                            "chunk_id": chunk_id,
                            "path": str(source_path),
                            "chunk_index": 0,
                        }
                    ],
                }
            ]
        },
    )

    load_calls = 0

    def _stub_load_config(_path: Path | None = None) -> AppConfig:
        nonlocal load_calls
        load_calls += 1
        return _build_config(runs_dir=runs_dir, markdown_dir=markdown_dir)

    monkeypatch.setattr("backend.api.routes.runs.load_config", _stub_load_config)

    app = create_app(
        runs_dir=runs_dir,
        max_workers=1,
        markdown_dir=markdown_dir,
        config_path=config_path,
    )
    with TestClient(app) as client:
        first = client.get(
            f"/api/v1/runs/run-source-chunks-cache/source-chunks?chunk_id={chunk_id}"
        )
        assert first.status_code == 200
        second = client.get(
            f"/api/v1/runs/run-source-chunks-cache/source-chunks?chunk_id={chunk_id}"
        )
        assert second.status_code == 200
        assert load_calls == 1

        updated_mtime_ns = config_path.stat().st_mtime_ns + 1_000_000_000
        os.utime(config_path, ns=(updated_mtime_ns, updated_mtime_ns))

        third = client.get(
            f"/api/v1/runs/run-source-chunks-cache/source-chunks?chunk_id={chunk_id}"
        )
        assert third.status_code == 200
        assert load_calls == 2


def test_api_duplicate_run_id_returns_conflict(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    runs_dir = tmp_path / "output"
    markdown_dir = tmp_path / "documents"
    markdown_dir.mkdir(parents=True, exist_ok=True)

    def _stub_load_config(_path: Path | None = None) -> AppConfig:
        return _build_config(runs_dir=runs_dir, markdown_dir=markdown_dir)

    def _stub_run_pipeline(
        question: str,
        config: AppConfig,
        run_id: str | None = None,
        log_llm_payload: bool = True,
        analysis_mode: str = "aggregate",
        api_key_override: str | None = None,
        selected_cities: list[str] | None = None,
    ) -> RunPaths:
        assert run_id is not None
        assert analysis_mode == "aggregate"
        assert api_key_override is None
        assert selected_cities is None
        return _write_success_artifacts(question=question, run_id=run_id, config=config)

    monkeypatch.setattr("backend.api.services.run_executor.load_config", _stub_load_config)
    monkeypatch.setattr("backend.api.services.run_executor.run_pipeline", _stub_run_pipeline)

    app = create_app(runs_dir=runs_dir, max_workers=1)
    with TestClient(app) as client:
        first = client.post(
            "/api/v1/runs",
            json={"question": "First", "run_id": "same-run"},
        )
        assert first.status_code == 202

        second = client.post(
            "/api/v1/runs",
            json={"question": "Second", "run_id": "same-run"},
        )
        assert second.status_code == 409


def test_api_run_id_is_not_blocked_by_stale_api_state_snapshot(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    runs_dir = tmp_path / "output"
    markdown_dir = tmp_path / "documents"
    markdown_dir.mkdir(parents=True, exist_ok=True)
    state_dir = runs_dir / "_api_state"
    state_dir.mkdir(parents=True, exist_ok=True)
    (state_dir / "stale-run.json").write_text(
        json.dumps(
            {
                "run_id": "stale-run",
                "question": "stale",
                "status": "failed",
                "started_at": datetime.now(timezone.utc).isoformat(),
                "completed_at": datetime.now(timezone.utc).isoformat(),
            },
            ensure_ascii=True,
            indent=2,
        ),
        encoding="utf-8",
    )

    def _stub_load_config(_path: Path | None = None) -> AppConfig:
        return _build_config(runs_dir=runs_dir, markdown_dir=markdown_dir)

    def _stub_run_pipeline(
        question: str,
        config: AppConfig,
        run_id: str | None = None,
        log_llm_payload: bool = True,
        analysis_mode: str = "aggregate",
        api_key_override: str | None = None,
        selected_cities: list[str] | None = None,
    ) -> RunPaths:
        assert run_id is not None
        assert analysis_mode == "aggregate"
        assert api_key_override is None
        assert selected_cities is None
        return _write_success_artifacts(question=question, run_id=run_id, config=config)

    monkeypatch.setattr("backend.api.services.run_executor.load_config", _stub_load_config)
    monkeypatch.setattr("backend.api.services.run_executor.run_pipeline", _stub_run_pipeline)

    app = create_app(runs_dir=runs_dir, max_workers=1, markdown_dir=markdown_dir)
    with TestClient(app) as client:
        start = client.post(
            "/api/v1/runs",
            json={"question": "reuse stale id", "run_id": "stale-run"},
        )
        assert start.status_code == 202
        terminal = _poll_until_terminal(client, "stale-run")
        assert terminal["status"] == "completed"


def test_api_status_not_found(tmp_path: Path) -> None:
    app = create_app(runs_dir=tmp_path / "output", max_workers=1)
    with TestClient(app) as client:
        response = client.get("/api/v1/runs/unknown/status")
        assert response.status_code == 404


def test_api_root_healthcheck(tmp_path: Path) -> None:
    app = create_app(runs_dir=tmp_path / "output", max_workers=1)
    with TestClient(app) as client:
        response = client.get("/")
        assert response.status_code == 200
        payload = response.json()
        assert payload["status"] == "ok"


def test_api_list_runs_reads_artifact_folders(tmp_path: Path) -> None:
    runs_dir = tmp_path / "output"
    markdown_dir = tmp_path / "documents"
    markdown_dir.mkdir(parents=True, exist_ok=True)
    config = _build_config(runs_dir=runs_dir, markdown_dir=markdown_dir)
    _write_success_artifacts(
        question="Historic run from artifact folder",
        run_id="run-from-folder",
        config=config,
    )

    app = create_app(runs_dir=runs_dir, max_workers=1, markdown_dir=markdown_dir)
    with TestClient(app) as client:
        response = client.get("/api/v1/runs")
        assert response.status_code == 200
        payload = response.json()
        assert payload["total"] == 1
        assert payload["runs"][0]["run_id"] == "run-from-folder"
        assert (
            payload["runs"][0]["question"] == "Historic run from artifact folder"
        )


def test_api_list_runs_reads_question_from_original_question_when_root_question_missing(
    tmp_path: Path,
) -> None:
    runs_dir = tmp_path / "output"
    markdown_dir = tmp_path / "documents"
    markdown_dir.mkdir(parents=True, exist_ok=True)

    run_dir = runs_dir / "run-inputs-question"
    run_dir.mkdir(parents=True, exist_ok=True)
    run_payload = {
        "run_id": "run-inputs-question",
        "inputs": {
            "original_question": "Question sourced from inputs.original_question",
            "canonical_research_query": "Primary fallback question",
        },
        "status": "completed",
        "started_at": datetime.now(timezone.utc).isoformat(),
        "completed_at": datetime.now(timezone.utc).isoformat(),
    }
    (run_dir / "run.json").write_text(
        json.dumps(run_payload, ensure_ascii=True, indent=2),
        encoding="utf-8",
    )

    app = create_app(runs_dir=runs_dir, max_workers=1, markdown_dir=markdown_dir)
    with TestClient(app) as client:
        response = client.get("/api/v1/runs")
        assert response.status_code == 200
        payload = response.json()
        assert payload["total"] == 1
        assert payload["runs"][0]["run_id"] == "run-inputs-question"
        assert (
            payload["runs"][0]["question"]
            == "Question sourced from inputs.original_question"
        )


def test_api_list_runs_reads_question_from_legacy_inputs_when_root_question_missing(
    tmp_path: Path,
) -> None:
    runs_dir = tmp_path / "output"
    markdown_dir = tmp_path / "documents"
    markdown_dir.mkdir(parents=True, exist_ok=True)

    run_dir = runs_dir / "run-legacy-inputs-question"
    run_dir.mkdir(parents=True, exist_ok=True)
    run_payload = {
        "run_id": "run-legacy-inputs-question",
        "inputs": {
            "initial_question": "Question sourced from inputs.initial_question",
        },
        "status": "completed",
        "started_at": datetime.now(timezone.utc).isoformat(),
        "completed_at": datetime.now(timezone.utc).isoformat(),
    }
    (run_dir / "run.json").write_text(
        json.dumps(run_payload, ensure_ascii=True, indent=2),
        encoding="utf-8",
    )

    app = create_app(runs_dir=runs_dir, max_workers=1, markdown_dir=markdown_dir)
    with TestClient(app) as client:
        response = client.get("/api/v1/runs")
        assert response.status_code == 200
        payload = response.json()
        assert payload["total"] == 1
        assert payload["runs"][0]["run_id"] == "run-legacy-inputs-question"
        assert (
            payload["runs"][0]["question"]
            == "Question sourced from inputs.initial_question"
        )

def test_api_list_runs_hides_failed_runs_by_default(tmp_path: Path) -> None:
    runs_dir = tmp_path / "output"
    markdown_dir = tmp_path / "documents"
    markdown_dir.mkdir(parents=True, exist_ok=True)
    config = _build_config(runs_dir=runs_dir, markdown_dir=markdown_dir)
    _write_run_listing_artifacts(
        question="Completed run",
        run_id="run-completed",
        status="completed",
        config=config,
        finish_reason="completed (write)",
    )
    _write_run_listing_artifacts(
        question="Failed run",
        run_id="run-failed",
        status="failed",
        config=config,
        finish_reason="writer_unexpected_error",
        error={"code": "RUN_EXECUTION_ERROR", "message": "Max turns (5) exceeded"},
    )

    app = create_app(runs_dir=runs_dir, max_workers=1, markdown_dir=markdown_dir)
    with TestClient(app) as client:
        response = client.get("/api/v1/runs")
        assert response.status_code == 200
        payload = response.json()
        assert payload["total"] == 1
        assert [item["run_id"] for item in payload["runs"]] == ["run-completed"]
        assert payload["runs"][0]["status"] == "completed"
        assert "picker_timestamp" in payload["runs"][0]


def test_api_list_runs_returns_picker_timestamp(tmp_path: Path) -> None:
    runs_dir = tmp_path / "output"
    markdown_dir = tmp_path / "documents"
    markdown_dir.mkdir(parents=True, exist_ok=True)
    _write_run_listing_artifact(
        runs_dir,
        run_id="run-picker-time",
        started_at=datetime(2026, 3, 12, 19, 54, tzinfo=timezone.utc),
        question="Timestamped picker run",
    )

    app = create_app(runs_dir=runs_dir, max_workers=1, markdown_dir=markdown_dir)
    with TestClient(app) as client:
        response = client.get("/api/v1/runs")
        assert response.status_code == 200
        payload = response.json()
        assert payload["total"] == 1
        assert payload["runs"][0]["picker_timestamp"] == "0312-1954"
        assert payload["runs"][0]["status"] == "completed"


def test_api_list_runs_include_all_returns_failed_runs_for_dev_mode(
    tmp_path: Path,
) -> None:
    runs_dir = tmp_path / "output"
    markdown_dir = tmp_path / "documents"
    markdown_dir.mkdir(parents=True, exist_ok=True)
    config = _build_config(runs_dir=runs_dir, markdown_dir=markdown_dir)
    _write_run_listing_artifacts(
        question="Completed run",
        run_id="run-completed",
        status="completed",
        config=config,
        finish_reason="completed (write)",
    )
    _write_run_listing_artifacts(
        question="Failed run",
        run_id="run-failed",
        status="failed",
        config=config,
        finish_reason="writer_unexpected_error",
        error={"code": "RUN_EXECUTION_ERROR", "message": "Max turns (5) exceeded"},
    )

    app = create_app(runs_dir=runs_dir, max_workers=1, markdown_dir=markdown_dir)
    with TestClient(app) as client:
        response = client.get("/api/v1/runs?include_all=true")
        assert response.status_code == 200
        payload = response.json()
        assert payload["total"] == 2
        assert {item["run_id"] for item in payload["runs"]} == {
            "run-completed",
            "run-failed",
        }
        failed_item = next(
            item for item in payload["runs"] if item["run_id"] == "run-failed"
        )
        assert failed_item["status"] == "failed"


def test_api_list_runs_search_matches_selected_city_with_typo_tolerance(
    tmp_path: Path,
) -> None:
    runs_dir = tmp_path / "output"
    markdown_dir = tmp_path / "documents"
    markdown_dir.mkdir(parents=True, exist_ok=True)
    _write_run_listing_artifact(
        runs_dir,
        run_id="run-leipzig",
        started_at=datetime(2026, 3, 12, 19, 54, tzinfo=timezone.utc),
        question="Charging rollout summary",
        inputs={
            "selected_cities_planned": [],
            "selected_cities_found": ["leipzig"],
        },
    )
    _write_run_listing_artifact(
        runs_dir,
        run_id="run-berlin",
        started_at=datetime(2026, 3, 13, 8, 30, tzinfo=timezone.utc),
        question="Charging rollout summary",
        inputs={
            "selected_cities_planned": [],
            "selected_cities_found": ["berlin"],
        },
    )

    app = create_app(runs_dir=runs_dir, max_workers=1, markdown_dir=markdown_dir)
    with TestClient(app) as client:
        response = client.get("/api/v1/runs?search=leipzing")
        assert response.status_code == 200
        payload = response.json()
        assert payload["total"] == 1
        assert [item["run_id"] for item in payload["runs"]] == ["run-leipzig"]


def test_api_list_runs_search_ranks_exact_question_phrase_before_token_match(
    tmp_path: Path,
) -> None:
    runs_dir = tmp_path / "output"
    markdown_dir = tmp_path / "documents"
    markdown_dir.mkdir(parents=True, exist_ok=True)
    _write_run_listing_artifact(
        runs_dir,
        run_id="run-token-match",
        started_at=datetime(2026, 3, 13, 9, 30, tzinfo=timezone.utc),
        question="Buses electric financing options",
    )
    _write_run_listing_artifact(
        runs_dir,
        run_id="run-exact-phrase",
        started_at=datetime(2026, 3, 12, 9, 30, tzinfo=timezone.utc),
        question="Electric buses financing options",
    )

    app = create_app(runs_dir=runs_dir, max_workers=1, markdown_dir=markdown_dir)
    with TestClient(app) as client:
        response = client.get("/api/v1/runs?search=electric+buses")
        assert response.status_code == 200
        payload = response.json()
        assert payload["total"] == 2
        assert [item["run_id"] for item in payload["runs"]] == [
            "run-exact-phrase",
            "run-token-match",
        ]


def test_api_list_runs_search_numeric_fragment_matches_run_id_or_question_only(
    tmp_path: Path,
) -> None:
    runs_dir = tmp_path / "output"
    markdown_dir = tmp_path / "documents"
    markdown_dir.mkdir(parents=True, exist_ok=True)
    _write_run_listing_artifact(
        runs_dir,
        run_id="20260326_1506",
        started_at=datetime(2026, 3, 26, 15, 6, tzinfo=timezone.utc),
        question="What initiatives exist for Munich?",
    )
    _write_run_listing_artifact(
        runs_dir,
        run_id="20260326_1511",
        started_at=datetime(2026, 3, 26, 15, 11, tzinfo=timezone.utc),
        question="What changed in project 1506 this quarter?",
    )
    _write_run_listing_artifact(
        runs_dir,
        run_id="gpt54mini-retrofit_rerun_dev-run3",
        started_at=datetime(2026, 3, 27, 9, 0, tzinfo=timezone.utc),
        question="What are the strongest retrofit initiatives in Munich?",
    )

    app = create_app(runs_dir=runs_dir, max_workers=1, markdown_dir=markdown_dir)
    with TestClient(app) as client:
        response = client.get("/api/v1/runs?search=1506")
        assert response.status_code == 200
        payload = response.json()
        assert payload["total"] == 2
        assert [item["run_id"] for item in payload["runs"]] == [
            "20260326_1506",
            "20260326_1511",
        ]


def test_api_run_diagnostics_returns_warning_and_error_artifacts(
    tmp_path: Path,
) -> None:
    """Diagnostics should expose run-local artifact labels plus parsed failure details."""
    runs_dir = tmp_path / "output"
    markdown_dir = tmp_path / "documents"
    markdown_dir.mkdir(parents=True, exist_ok=True)
    config = _build_config(runs_dir=runs_dir, markdown_dir=markdown_dir)
    paths = _write_run_listing_artifacts(
        question="Failed run",
        run_id="run-failed",
        status="failed",
        config=config,
        finish_reason="writer_unexpected_error",
        error={"code": "RUN_EXECUTION_ERROR", "message": "Max turns (5) exceeded"},
    )
    run_payload = json.loads(paths.run_log.read_text(encoding="utf-8"))
    run_payload["artifacts"] = {
        "run_summary": str(paths.run_summary),
        "error_log": str(paths.error_log),
    }
    run_payload["llm_usage"] = {"calls": 2, "totals": {"total_tokens": 123}}
    run_payload["retry_summary"] = {"total_events": 1, "by_operation": {"writer": 1}}
    run_payload["writer_multi_pass"] = {
        "strategy": "split_by_city",
        "combine_strategy": "draft_merge",
        "analysis_mode": "aggregate",
        "payload_tokens": 350083,
        "threshold_tokens": 200000,
        "batch_count": 2,
        "batches": [
            {
                "batch_index": 1,
                "city_names": ["Aachen", "Amsterdam"],
                "excerpt_count": 1300,
                "payload_tokens": 175040,
            },
            {
                "batch_index": 2,
                "city_names": ["Antwerp", "Athens"],
                "excerpt_count": 1274,
                "payload_tokens": 174980,
            },
        ],
    }
    run_payload["writer_section_plan"] = {
        "strategy": "section_first",
        "analysis_mode": "aggregate",
        "planner_input_tokens": 2400,
        "catalog_truncated": False,
        "section_count": 1,
        "sections": [
            {
                "section_id": "retrofit_investment",
                "title": "Retrofit Investment Evidence",
                "section_type": "numeric_analysis",
                "purpose": "Compare retrofit investment evidence.",
                "required_ref_ids": ["ref_1", "ref_2"],
                "city_names": ["Aachen", "Amsterdam"],
                "writing_instructions": "Compare the assigned evidence.",
                "payload_tokens": 1800,
                "draft_length_chars": 420,
                "batch_count": 1,
            }
        ],
    }
    paths.run_log.write_text(
        json.dumps(run_payload, ensure_ascii=True, indent=2),
        encoding="utf-8",
    )
    paths.error_log.write_text(
        "\n".join(
            [
                "2026-01-01 00:00:01 worker.py:11 - ERROR - writer crashed",
                "Traceback line",
            ]
        ),
        encoding="utf-8",
    )
    paths.run_summary.write_text("RUN SUMMARY\nStatus: failed", encoding="utf-8")
    (paths.base_dir / "run.log").write_text(
        "\n".join(
            [
                "2026-01-01 00:00:00 worker.py:10 - INFO - setup",
                (
                    '2026-01-01 00:00:00 worker.py:10 - WARNING - WRITER_CITATION_COVERAGE '
                    '{"run_id":"run-failed","attempt":2,"max_attempts":2,"status":"exhausted",'
                    '"coverage_confirmed":57,"coverage_required":101,"coverage_ratio":"57/101",'
                    '"missing_cities":["Antwerp","Bergamo"],"analysis_mode":"aggregate"}'
                ),
                "2026-01-01 00:00:01 worker.py:11 - WARNING - writer nearing max turns",
                "2026-01-01 00:00:02 worker.py:12 - ERROR - writer crashed",
                "Traceback line",
                "2026-01-01 00:00:03 worker.py:13 - INFO - done",
            ]
        ),
        encoding="utf-8",
    )

    app = create_app(runs_dir=runs_dir, max_workers=1, markdown_dir=markdown_dir)
    with TestClient(app) as client:
        response = client.get("/api/v1/runs/run-failed/diagnostics")
        assert response.status_code == 200
        payload = response.json()
        assert payload["run_id"] == "run-failed"
        assert payload["status"] == "failed"
        assert payload["error"]["code"] == "RUN_EXECUTION_ERROR"
        assert payload["writer_citation_coverage"]["status"] == "exhausted"
        assert payload["writer_citation_coverage"]["coverage_ratio"] == "57/101"
        assert payload["writer_citation_coverage"]["missing_cities"] == [
            "Antwerp",
            "Bergamo",
        ]
        assert payload["writer_multi_pass"]["batch_count"] == 2
        assert payload["writer_multi_pass"]["threshold_tokens"] == 200000
        assert payload["writer_multi_pass"]["batches"][0]["city_names"] == [
            "Aachen",
            "Amsterdam",
        ]
        assert payload["writer_section_plan"]["strategy"] == "section_first"
        assert payload["writer_section_plan"]["section_count"] == 1
        assert payload["writer_section_plan"]["sections"][0]["section_id"] == (
            "retrofit_investment"
        )
        assert payload["warning_entries"] == [
            (
                '2026-01-01 00:00:00 worker.py:10 - WARNING - WRITER_CITATION_COVERAGE '
                '{"run_id":"run-failed","attempt":2,"max_attempts":2,"status":"exhausted",'
                '"coverage_confirmed":57,"coverage_required":101,"coverage_ratio":"57/101",'
                '"missing_cities":["Antwerp","Bergamo"],"analysis_mode":"aggregate"}'
            ),
            "2026-01-01 00:00:01 worker.py:11 - WARNING - writer nearing max turns",
        ]
        assert payload["error_log_text"].startswith(
            "2026-01-01 00:00:01 worker.py:11 - ERROR - writer crashed"
        )
        assert payload["artifacts"]["run_summary"] == "run_summary.txt"
        assert payload["artifacts"]["error_log"] == "error_log.txt"
        assert payload["artifacts"]["run_log"] == "run.log"
        assert payload["retry_summary"]["total_events"] == 1
        assert payload["llm_usage"]["calls"] == 2


def test_api_run_diagnostics_ignores_foreign_artifact_paths(tmp_path: Path) -> None:
    """Diagnostics should ignore foreign artifact paths and stay inside the run folder."""
    runs_dir = tmp_path / "output"
    markdown_dir = tmp_path / "documents"
    markdown_dir.mkdir(parents=True, exist_ok=True)
    config = _build_config(runs_dir=runs_dir, markdown_dir=markdown_dir)
    paths = _write_run_listing_artifacts(
        question="Failed run",
        run_id="run-foreign-artifacts",
        status="failed",
        config=config,
        finish_reason="writer_unexpected_error",
        error={"code": "RUN_EXECUTION_ERROR", "message": "Max turns (5) exceeded"},
    )

    foreign_dir = tmp_path / "foreign-artifacts"
    foreign_dir.mkdir(parents=True, exist_ok=True)
    foreign_summary = foreign_dir / "run_summary.txt"
    foreign_error_log = foreign_dir / "error_log.txt"
    foreign_summary.write_text("FOREIGN SUMMARY", encoding="utf-8")
    foreign_error_log.write_text("FOREIGN SECRET", encoding="utf-8")

    run_payload = json.loads(paths.run_log.read_text(encoding="utf-8"))
    run_payload["artifacts"] = {
        "run_summary": str(foreign_summary),
        "error_log": str(foreign_error_log),
    }
    paths.run_log.write_text(
        json.dumps(run_payload, ensure_ascii=True, indent=2),
        encoding="utf-8",
    )
    paths.run_summary.write_text("LOCAL SUMMARY", encoding="utf-8")
    paths.error_log.write_text("LOCAL ERROR", encoding="utf-8")
    (paths.base_dir / "run.log").write_text(
        "\n".join(
            [
                "2026-01-01 00:00:00 worker.py:10 - INFO - setup",
                "2026-01-01 00:00:01 worker.py:11 - ERROR - local failure",
            ]
        ),
        encoding="utf-8",
    )

    app = create_app(runs_dir=runs_dir, max_workers=1, markdown_dir=markdown_dir)
    with TestClient(app) as client:
        response = client.get("/api/v1/runs/run-foreign-artifacts/diagnostics")
        assert response.status_code == 200
        payload = response.json()
        assert payload["artifacts"]["run_summary"] == "run_summary.txt"
        assert payload["artifacts"]["error_log"] == "error_log.txt"
        assert payload["artifacts"]["run_log"] == "run.log"
        assert payload["error_log_text"] == "LOCAL ERROR"
        assert "FOREIGN SECRET" not in payload["error_log_text"]


def test_api_output_and_context_resolve_stale_container_artifact_paths(
    tmp_path: Path,
) -> None:
    runs_dir = tmp_path / "output"
    markdown_dir = tmp_path / "documents"
    markdown_dir.mkdir(parents=True, exist_ok=True)
    config = _build_config(runs_dir=runs_dir, markdown_dir=markdown_dir)
    run_id = "run-stale-paths"
    paths = _write_success_artifacts(
        question="Historic run with stale artifact paths",
        run_id=run_id,
        config=config,
    )
    run_payload = json.loads(paths.run_log.read_text(encoding="utf-8"))
    run_payload["artifacts"]["final_output"] = f"/data/output/{run_id}/final.md"
    run_payload["artifacts"]["context_bundle"] = f"/data/output/{run_id}/context_bundle.json"
    paths.run_log.write_text(
        json.dumps(run_payload, ensure_ascii=True, indent=2),
        encoding="utf-8",
    )

    app = create_app(runs_dir=runs_dir, max_workers=1, markdown_dir=markdown_dir)
    with TestClient(app) as client:
        output_response = client.get(f"/api/v1/runs/{run_id}/output")
        assert output_response.status_code == 200
        assert "Stub answer" in output_response.json()["content"]

        context_response = client.get(f"/api/v1/runs/{run_id}/context")
        assert context_response.status_code == 200
        assert isinstance(context_response.json()["context_bundle"], dict)


def test_api_docx_export_returns_word_document(tmp_path: Path) -> None:
    runs_dir = tmp_path / "output"
    markdown_dir = tmp_path / "documents"
    markdown_dir.mkdir(parents=True, exist_ok=True)
    config = _build_config(runs_dir=runs_dir, markdown_dir=markdown_dir)
    run_id = "run-docx-export"
    paths = _write_success_artifacts(
        question="Export run",
        run_id=run_id,
        config=config,
    )
    paths.final_output.write_text(
        "# Export report\n\n"
        "Munich remains on schedule. [ref_1][ref_2]\n\n"
        "| City | Comment mode |\n"
        "| --- | --- |\n"
        "| Munich | Google Doc review [ref_3] |\n",
        encoding="utf-8",
    )

    app = create_app(runs_dir=runs_dir, max_workers=1, markdown_dir=markdown_dir)
    with TestClient(app) as client:
        response = client.get(f"/api/v1/runs/{run_id}/export/docx")
        assert response.status_code == 200
        assert (
            response.headers["content-type"]
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        assert f'filename="{run_id}.docx"' in response.headers["content-disposition"]

    document = Document(BytesIO(response.content))
    assert document.paragraphs[0].text == "Export report"
    assert document.paragraphs[1].text == "Munich remains on schedule."
    assert len(document.tables) == 1
    assert document.tables[0].rows[1].cells[0].text == "Munich"
    assert document.tables[0].rows[1].cells[1].text == "Google Doc review"


def test_api_writer_context_export_returns_writer_safe_json_bundle(tmp_path: Path) -> None:
    runs_dir = tmp_path / "output"
    markdown_dir = tmp_path / "documents"
    markdown_dir.mkdir(parents=True, exist_ok=True)
    config = _build_config(runs_dir=runs_dir, markdown_dir=markdown_dir)
    run_id = "run-writer-context-export"
    paths = _write_success_artifacts(
        question="Writer context export run",
        run_id=run_id,
        config=config,
    )
    context_bundle = {
        "research_question": "What retrofit evidence was selected?",
        "analysis_mode": "aggregate",
        "selected_cities": ["munich", "leipzig"],
        "enrichment": {"status": "success", "notes": ["not sent to writer"]},
        "final": str(paths.final_output),
        "markdown": {
            "status": "success",
            "analysis_mode": "aggregate",
            "selected_city_names": ["Munich", "Leipzig"],
            "inspected_city_names": ["Munich", "Leipzig"],
            "selected_cities": ["munich", "leipzig"],
            "inspected_cities": ["munich", "leipzig"],
            "accepted_chunk_ids": ["chunk_munich_1", "chunk_leipzig_1"],
            "rejected_chunk_ids": ["chunk_munich_9"],
            "decision_audit": {"accepted_total": 2, "rejected_total": 1},
            "excerpt_count": 2,
            "excerpts": [
                {
                    "ref_id": "ref_1",
                    "city_name": "Munich",
                    "city_key": "munich",
                    "quote": "Munich is retrofitting schools.",
                    "partial_answer": "Munich is retrofitting schools.",
                    "source_chunk_ids": ["chunk_munich_1"],
                },
                {
                    "ref_id": "ref_2",
                    "city_name": "Leipzig",
                    "city_key": "leipzig",
                    "quote": "Leipzig expanded district heating.",
                    "partial_answer": "Leipzig expanded district heating.",
                    "source_chunk_ids": ["chunk_leipzig_1"],
                },
            ],
        },
    }
    paths.context_bundle.write_text(
        json.dumps(context_bundle, ensure_ascii=True, indent=2),
        encoding="utf-8",
    )

    app = create_app(runs_dir=runs_dir, max_workers=1, markdown_dir=markdown_dir)
    with TestClient(app) as client:
        response = client.get(f"/api/v1/runs/{run_id}/export/writer-context")
        assert response.status_code == 200
        assert response.headers["content-type"].startswith("application/json")
        assert (
            f'filename="{run_id}_writer_context.json"'
            in response.headers["content-disposition"]
        )
        payload = response.json()

    assert payload["research_question"] == "What retrofit evidence was selected?"
    assert payload["analysis_mode"] == "aggregate"
    assert payload["selected_cities"] == ["Munich", "Leipzig"]
    assert "sql" not in payload
    markdown_payload = payload["markdown"]
    assert markdown_payload["excerpt_count"] == 2
    assert markdown_payload["excerpts"][0]["ref_id"] == "ref_1"
    assert markdown_payload["excerpts"][0]["source_chunk_ids"] == ["chunk_munich_1"]
    assert "enrichment" not in payload
    assert "accepted_chunk_ids" not in markdown_payload
    assert "decision_audit" not in markdown_payload


def test_api_writer_context_markdown_export_remains_available(tmp_path: Path) -> None:
    runs_dir = tmp_path / "output"
    markdown_dir = tmp_path / "documents"
    markdown_dir.mkdir(parents=True, exist_ok=True)
    config = _build_config(runs_dir=runs_dir, markdown_dir=markdown_dir)
    run_id = "run-writer-context-markdown-export"
    paths = _write_success_artifacts(
        question="Writer context markdown export run",
        run_id=run_id,
        config=config,
    )
    context_bundle = {
        "research_question": "What retrofit evidence was selected?",
        "analysis_mode": "aggregate",
        "selected_cities": ["munich", "leipzig"],
        "enrichment": {"status": "success", "notes": ["not sent to writer"]},
        "final": str(paths.final_output),
        "markdown": {
            "status": "success",
            "analysis_mode": "aggregate",
            "selected_city_names": ["Munich", "Leipzig"],
            "inspected_city_names": ["Munich", "Leipzig"],
            "selected_cities": ["munich", "leipzig"],
            "inspected_cities": ["munich", "leipzig"],
            "accepted_chunk_ids": ["chunk_munich_1", "chunk_leipzig_1"],
            "decision_audit": {"accepted_total": 2, "rejected_total": 1},
            "excerpt_count": 1,
            "excerpts": [
                {
                    "ref_id": "ref_1",
                    "city_name": "Munich",
                    "city_key": "munich",
                    "quote": "Munich is retrofitting schools.",
                    "partial_answer": "Munich is retrofitting schools.",
                    "source_chunk_ids": ["chunk_munich_1"],
                },
            ],
        },
    }
    paths.context_bundle.write_text(
        json.dumps(context_bundle, ensure_ascii=True, indent=2),
        encoding="utf-8",
    )

    app = create_app(runs_dir=runs_dir, max_workers=1, markdown_dir=markdown_dir)
    with TestClient(app) as client:
        response = client.get(f"/api/v1/runs/{run_id}/export/writer-context.md")
        assert response.status_code == 200
        assert response.headers["content-type"].startswith("text/markdown")
        assert (
            f'filename="{run_id}_writer_context.md"'
            in response.headers["content-disposition"]
        )
        payload = response.text

    assert "# Writer Context Export" in payload
    assert "- Research question: What retrofit evidence was selected?" in payload
    assert "- Selected cities: Munich, Leipzig" in payload
    assert "## SQL Context" not in payload
    assert "## Excerpt 1 - Munich (`ref_1`)" in payload
    assert "> Munich is retrofitting schools." in payload
    assert "not sent to writer" not in payload
    assert "accepted_chunk_ids" not in payload
    assert "decision_audit" not in payload


def test_api_output_hides_legacy_finish_reason_footer(tmp_path: Path) -> None:
    runs_dir = tmp_path / "output"
    markdown_dir = tmp_path / "documents"
    markdown_dir.mkdir(parents=True, exist_ok=True)
    config = _build_config(runs_dir=runs_dir, markdown_dir=markdown_dir)
    run_id = "run-output-legacy-footer"
    paths = _write_success_artifacts(
        question="Legacy footer run",
        run_id=run_id,
        config=config,
    )
    paths.final_output.write_text(
        "# Question\nLegacy footer run\n\n"
        "# Answer\nVisible answer body.\n\n"
        "---\n"
        "Finish reason: completed (write)\n",
        encoding="utf-8",
    )

    app = create_app(runs_dir=runs_dir, max_workers=1, markdown_dir=markdown_dir)
    with TestClient(app) as client:
        response = client.get(f"/api/v1/runs/{run_id}/output")
        assert response.status_code == 200
        payload = response.json()
        assert payload["content"] == "# Question\nLegacy footer run\n\n# Answer\nVisible answer body."
        assert "Finish reason:" not in payload["content"]


def test_api_list_runs_drops_entry_after_artifact_folder_is_deleted(tmp_path: Path) -> None:
    runs_dir = tmp_path / "output"
    markdown_dir = tmp_path / "documents"
    markdown_dir.mkdir(parents=True, exist_ok=True)
    config = _build_config(runs_dir=runs_dir, markdown_dir=markdown_dir)
    paths = _write_success_artifacts(
        question="Run to be deleted from disk",
        run_id="run-deleted",
        config=config,
    )

    app = create_app(runs_dir=runs_dir, max_workers=1, markdown_dir=markdown_dir)
    with TestClient(app) as client:
        before_delete = client.get("/api/v1/runs")
        assert before_delete.status_code == 200
        before_payload = before_delete.json()
        assert before_payload["total"] == 1
        assert before_payload["runs"][0]["run_id"] == "run-deleted"

        shutil.rmtree(paths.base_dir)

        after_delete = client.get("/api/v1/runs")
        assert after_delete.status_code == 200
        after_payload = after_delete.json()
        assert after_payload["total"] == 0


def test_api_output_returns_conflict_while_running(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    runs_dir = tmp_path / "output"
    markdown_dir = tmp_path / "documents"
    markdown_dir.mkdir(parents=True, exist_ok=True)
    started = threading.Event()
    release = threading.Event()

    def _stub_load_config(_path: Path | None = None) -> AppConfig:
        return _build_config(runs_dir=runs_dir, markdown_dir=markdown_dir)

    def _slow_run_pipeline(
        question: str,
        config: AppConfig,
        run_id: str | None = None,
        log_llm_payload: bool = True,
        analysis_mode: str = "aggregate",
        api_key_override: str | None = None,
        selected_cities: list[str] | None = None,
    ) -> RunPaths:
        assert run_id is not None
        assert analysis_mode == "aggregate"
        assert api_key_override is None
        assert selected_cities is None
        started.set()
        release.wait(timeout=2)
        return _write_success_artifacts(question=question, run_id=run_id, config=config)

    monkeypatch.setattr("backend.api.services.run_executor.load_config", _stub_load_config)
    monkeypatch.setattr("backend.api.services.run_executor.run_pipeline", _slow_run_pipeline)

    app = create_app(runs_dir=runs_dir, max_workers=1)
    with TestClient(app) as client:
        start = client.post(
            "/api/v1/runs",
            json={"question": "Running test", "run_id": "run-running"},
        )
        assert start.status_code == 202
        assert started.wait(timeout=1)

        output_response = client.get("/api/v1/runs/run-running/output")
        assert output_response.status_code == 409

        release.set()
        terminal = _poll_until_terminal(client, "run-running")
        assert terminal["status"] == "completed"


def test_api_failed_run(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    runs_dir = tmp_path / "output"
    markdown_dir = tmp_path / "documents"
    markdown_dir.mkdir(parents=True, exist_ok=True)

    def _stub_load_config(_path: Path | None = None) -> AppConfig:
        return _build_config(runs_dir=runs_dir, markdown_dir=markdown_dir)

    def _failing_run_pipeline(
        question: str,
        config: AppConfig,
        run_id: str | None = None,
        log_llm_payload: bool = True,
        analysis_mode: str = "aggregate",
        api_key_override: str | None = None,
        selected_cities: list[str] | None = None,
    ) -> RunPaths:
        assert analysis_mode == "aggregate"
        assert api_key_override is None
        assert selected_cities is None
        raise RuntimeError("simulated pipeline failure")

    monkeypatch.setattr("backend.api.services.run_executor.load_config", _stub_load_config)
    monkeypatch.setattr(
        "backend.api.services.run_executor.run_pipeline", _failing_run_pipeline
    )

    app = create_app(runs_dir=runs_dir, max_workers=1)
    with TestClient(app) as client:
        start = client.post(
            "/api/v1/runs",
            json={"question": "Failing test", "run_id": "run-failed"},
        )
        assert start.status_code == 202

        terminal = _poll_until_terminal(client, "run-failed")
        assert terminal["status"] == "failed"
        assert terminal["error"]["code"] == "RUN_EXECUTION_ERROR"

        output_response = client.get("/api/v1/runs/run-failed/output")
        assert output_response.status_code == 409

        context_response = client.get("/api/v1/runs/run-failed/context")
        assert context_response.status_code == 409


def test_api_failed_run_uses_persisted_decision_error(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    runs_dir = tmp_path / "output"
    markdown_dir = tmp_path / "documents"
    markdown_dir.mkdir(parents=True, exist_ok=True)
    finish_reason = "query_preparation_failed"
    error_code = "QUERY_PREPARATION_ERROR"
    error_message = (
        "Could not prepare the research query for this request. Please try again."
    )

    def _stub_load_config(_path: Path | None = None) -> AppConfig:
        return _build_config(runs_dir=runs_dir, markdown_dir=markdown_dir)

    def _failed_run_pipeline(
        question: str,
        config: AppConfig,
        run_id: str | None = None,
        log_llm_payload: bool = True,
        analysis_mode: str = "aggregate",
        api_key_override: str | None = None,
        selected_cities: list[str] | None = None,
    ) -> RunPaths:
        assert run_id is not None
        assert isinstance(log_llm_payload, bool)
        assert analysis_mode == "aggregate"
        assert api_key_override is None
        assert selected_cities is None
        return _write_failed_artifacts_with_decision_error(
            question=question,
            run_id=run_id,
            config=config,
            finish_reason=finish_reason,
            error_code=error_code,
            error_message=error_message,
        )

    monkeypatch.setattr("backend.api.services.run_executor.load_config", _stub_load_config)
    monkeypatch.setattr(
        "backend.api.services.run_executor.run_pipeline", _failed_run_pipeline
    )

    app = create_app(runs_dir=runs_dir, max_workers=1, markdown_dir=markdown_dir)
    with TestClient(app) as client:
        start = client.post(
            "/api/v1/runs",
            json={
                "question": "Query preparation failed",
                "run_id": "run-persisted-failure",
            },
        )
        assert start.status_code == 202
        terminal = _poll_until_terminal(client, "run-persisted-failure")
        assert terminal["status"] == "failed"
        assert terminal["finish_reason"] == finish_reason
        assert terminal["error"]["code"] == error_code
        assert terminal["error"]["message"] == error_message


def test_api_failed_run_preserves_persisted_failure_details_after_exception(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    runs_dir = tmp_path / "output"
    markdown_dir = tmp_path / "documents"
    markdown_dir.mkdir(parents=True, exist_ok=True)
    finish_reason = "query_preparation_failed"
    error_code = "QUERY_PREPARATION_ERROR"
    error_message = (
        "Could not prepare the research query for this request. Please try again."
    )

    def _stub_load_config(_path: Path | None = None) -> AppConfig:
        return _build_config(runs_dir=runs_dir, markdown_dir=markdown_dir)

    def _failing_run_pipeline(
        question: str,
        config: AppConfig,
        run_id: str | None = None,
        log_llm_payload: bool = True,
        analysis_mode: str = "aggregate",
        api_key_override: str | None = None,
        selected_cities: list[str] | None = None,
    ) -> RunPaths:
        assert run_id is not None
        assert isinstance(log_llm_payload, bool)
        assert analysis_mode == "aggregate"
        assert api_key_override is None
        assert selected_cities is None
        _write_failed_artifacts_with_decision_error(
            question=question,
            run_id=run_id,
            config=config,
            finish_reason=finish_reason,
            error_code=error_code,
            error_message=error_message,
        )
        raise ValueError("generic executor wrapper failure")

    monkeypatch.setattr("backend.api.services.run_executor.load_config", _stub_load_config)
    monkeypatch.setattr(
        "backend.api.services.run_executor.run_pipeline", _failing_run_pipeline
    )

    app = create_app(runs_dir=runs_dir, max_workers=1, markdown_dir=markdown_dir)
    with TestClient(app) as client:
        start = client.post(
            "/api/v1/runs",
            json={
                "question": "Query preparation failed with exception",
                "run_id": "run-preserved-failure",
            },
        )
        assert start.status_code == 202
        terminal = _poll_until_terminal(client, "run-preserved-failure")
        assert terminal["status"] == "failed"
        assert terminal["finish_reason"] == finish_reason
        assert terminal["error"]["code"] == error_code
        assert terminal["error"]["message"] == error_message

    run_log_payload = json.loads(
        (runs_dir / "run-preserved-failure" / "run.json").read_text(encoding="utf-8")
    )
    assert run_log_payload["finish_reason"] == finish_reason
    assert run_log_payload["error"]["code"] == error_code
    assert run_log_payload["error"]["message"] == error_message


def test_api_failed_run_writes_error_log_snapshot_for_executor_failures(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    runs_dir = tmp_path / "output"
    markdown_dir = tmp_path / "documents"
    markdown_dir.mkdir(parents=True, exist_ok=True)

    def _stub_load_config(_path: Path | None = None) -> AppConfig:
        return _build_config(runs_dir=runs_dir, markdown_dir=markdown_dir)

    def _failing_run_pipeline(
        question: str,
        config: AppConfig,
        run_id: str | None = None,
        log_llm_payload: bool = True,
        analysis_mode: str = "aggregate",
        api_key_override: str | None = None,
        selected_cities: list[str] | None = None,
    ) -> RunPaths:
        assert run_id is not None
        assert analysis_mode == "aggregate"
        assert api_key_override is None
        assert selected_cities is None
        _write_started_artifacts_with_error_log_input(
            question=question,
            run_id=run_id,
            config=config,
        )
        raise RuntimeError("simulated pipeline failure")

    monkeypatch.setattr("backend.api.services.run_executor.load_config", _stub_load_config)
    monkeypatch.setattr(
        "backend.api.services.run_executor.run_pipeline", _failing_run_pipeline
    )

    app = create_app(runs_dir=runs_dir, max_workers=1)
    with TestClient(app) as client:
        start = client.post(
            "/api/v1/runs",
            json={
                "question": "Failure should still emit error_log",
                "run_id": "run-failed-with-error-log",
            },
        )
        assert start.status_code == 202
        terminal = _poll_until_terminal(client, "run-failed-with-error-log")
        assert terminal["status"] == "failed"
        assert terminal["error"]["code"] == "RUN_EXECUTION_ERROR"

    run_dir = runs_dir / "run-failed-with-error-log"
    error_log_path = run_dir / "error_log.txt"
    assert error_log_path.exists()
    error_lines = error_log_path.read_text(encoding="utf-8")
    assert " - ERROR - writer crashed" in error_lines
    assert " - CRITICAL - aborting run" in error_lines

    run_log_payload = json.loads((run_dir / "run.json").read_text(encoding="utf-8"))
    assert run_log_payload["status"] == "failed"
    assert run_log_payload["finish_reason"] == "run_execution_error"
    assert run_log_payload["error"]["code"] == "RUN_EXECUTION_ERROR"
    assert run_log_payload["artifacts"]["error_log"] == str(error_log_path)


def test_api_run_filters_markdown_by_selected_cities(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    runs_dir = tmp_path / "output"
    markdown_dir = tmp_path / "documents"
    markdown_dir.mkdir(parents=True, exist_ok=True)
    (markdown_dir / "Berlin.md").write_text("# Berlin", encoding="utf-8")
    (markdown_dir / "Munich.md").write_text("# Munich", encoding="utf-8")
    captured_files: list[str] = []

    def _stub_load_config(_path: Path | None = None) -> AppConfig:
        return _build_config(runs_dir=runs_dir, markdown_dir=markdown_dir)

    def _stub_run_pipeline(
        question: str,
        config: AppConfig,
        run_id: str | None = None,
        log_llm_payload: bool = True,
        analysis_mode: str = "aggregate",
        api_key_override: str | None = None,
        selected_cities: list[str] | None = None,
    ) -> RunPaths:
        assert run_id is not None
        assert analysis_mode == "aggregate"
        assert api_key_override is None
        assert selected_cities == ["Berlin"]
        captured_files.extend(
            sorted(path.name for path in config.markdown_dir.rglob("*.md"))
        )
        return _write_success_artifacts(question=question, run_id=run_id, config=config)

    monkeypatch.setattr("backend.api.services.run_executor.load_config", _stub_load_config)
    monkeypatch.setattr("backend.api.services.run_executor.run_pipeline", _stub_run_pipeline)

    app = create_app(runs_dir=runs_dir, max_workers=1, markdown_dir=markdown_dir)
    with TestClient(app) as client:
        start = client.post(
            "/api/v1/runs",
            json={
                "question": "Only Berlin please",
                "run_id": "run-berlin",
                "cities": ["Berlin"],
            },
        )
        assert start.status_code == 202
        terminal = _poll_until_terminal(client, "run-berlin")
        assert terminal["status"] == "completed"
        listed_runs = client.get("/api/v1/runs")
        assert listed_runs.status_code == 200
        listed_ids = [item["run_id"] for item in listed_runs.json()["runs"]]
        assert listed_ids == ["run-berlin"]

    assert captured_files == ["Berlin.md"]


def test_api_run_preserves_display_city_names_while_deduping_aliases(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    runs_dir = tmp_path / "output"
    markdown_dir = tmp_path / "documents"
    markdown_dir.mkdir(parents=True, exist_ok=True)
    (markdown_dir / "Istanbul.md").write_text("# Istanbul", encoding="utf-8")

    def _stub_load_config(_path: Path | None = None) -> AppConfig:
        return _build_config(runs_dir=runs_dir, markdown_dir=markdown_dir)

    def _stub_run_pipeline(
        question: str,
        config: AppConfig,
        run_id: str | None = None,
        log_llm_payload: bool = True,
        analysis_mode: str = "aggregate",
        api_key_override: str | None = None,
        selected_cities: list[str] | None = None,
    ) -> RunPaths:
        assert run_id is not None
        assert isinstance(log_llm_payload, bool)
        assert analysis_mode == "aggregate"
        assert api_key_override is None
        assert selected_cities == ["Istanbul"]
        return _write_success_artifacts(question=question, run_id=run_id, config=config)

    monkeypatch.setattr("backend.api.services.run_executor.load_config", _stub_load_config)
    monkeypatch.setattr("backend.api.services.run_executor.run_pipeline", _stub_run_pipeline)

    app = create_app(runs_dir=runs_dir, max_workers=1, markdown_dir=markdown_dir)
    with TestClient(app) as client:
        start = client.post(
            "/api/v1/runs",
            json={
                "question": "Only Istanbul please",
                "run_id": "run-istanbul",
                "cities": ["Istanbul", "ISTANBUL"],
            },
        )
        assert start.status_code == 202
        terminal = _poll_until_terminal(client, "run-istanbul")
        assert terminal["status"] == "completed"


def test_api_run_analysis_mode_defaults_and_passes_explicit_value(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    runs_dir = tmp_path / "output"
    markdown_dir = tmp_path / "documents"
    markdown_dir.mkdir(parents=True, exist_ok=True)
    captured_modes: list[str] = []

    def _stub_load_config(_path: Path | None = None) -> AppConfig:
        return _build_config(runs_dir=runs_dir, markdown_dir=markdown_dir)

    def _stub_run_pipeline(
        question: str,
        config: AppConfig,
        run_id: str | None = None,
        log_llm_payload: bool = True,
        analysis_mode: str = "aggregate",
        api_key_override: str | None = None,
        selected_cities: list[str] | None = None,
    ) -> RunPaths:
        assert run_id is not None
        assert selected_cities is None
        assert api_key_override is None
        captured_modes.append(analysis_mode)
        return _write_success_artifacts(question=question, run_id=run_id, config=config)

    monkeypatch.setattr("backend.api.services.run_executor.load_config", _stub_load_config)
    monkeypatch.setattr("backend.api.services.run_executor.run_pipeline", _stub_run_pipeline)

    app = create_app(runs_dir=runs_dir, max_workers=1, markdown_dir=markdown_dir)
    with TestClient(app) as client:
        default_start = client.post(
            "/api/v1/runs",
            json={"question": "Default mode", "run_id": "run-mode-default"},
        )
        assert default_start.status_code == 202
        _poll_until_terminal(client, "run-mode-default")

        explicit_start = client.post(
            "/api/v1/runs",
            json={
                "question": "City mode",
                "run_id": "run-mode-city-by-city",
                "analysis_mode": "city_by_city",
            },
        )
        assert explicit_start.status_code == 202
        _poll_until_terminal(client, "run-mode-city-by-city")

    assert captured_modes == ["aggregate", "city_by_city"]


def test_api_run_rejects_invalid_analysis_mode(tmp_path: Path) -> None:
    app = create_app(runs_dir=tmp_path / "output", max_workers=1)
    with TestClient(app) as client:
        response = client.post(
            "/api/v1/runs",
            json={
                "question": "Invalid mode",
                "run_id": "run-invalid-mode",
                "analysis_mode": "invalid_mode",
            },
        )
        assert response.status_code == 422


def test_api_list_runs_ignores_stale_api_state_and_uses_artifact_folder(tmp_path: Path) -> None:
    runs_dir = tmp_path / "output"
    markdown_dir = tmp_path / "documents"
    markdown_dir.mkdir(parents=True, exist_ok=True)

    state_dir = runs_dir / "_api_state"
    state_dir.mkdir(parents=True, exist_ok=True)
    legacy_run_dir = runs_dir / "legacy-run_01"
    legacy_run_dir.mkdir(parents=True, exist_ok=True)
    run_log_path = legacy_run_dir / "run.json"
    started_at = datetime.now(timezone.utc).isoformat()
    completed_at = datetime.now(timezone.utc).isoformat()

    run_log_path.write_text(
        json.dumps(
            {
                "run_id": "legacy-run_01",
                "question": "Legacy alias run",
                "status": "completed",
                "started_at": started_at,
                "completed_at": completed_at,
                "finish_reason": "completed (write)",
                "artifacts": {
                    "context_bundle": str(legacy_run_dir / "context_bundle.json"),
                    "final_output": str(legacy_run_dir / "final.md"),
                },
            },
            ensure_ascii=True,
            indent=2,
        ),
        encoding="utf-8",
    )

    shared_payload = {
        "question": "Legacy alias run",
        "status": "completed",
        "started_at": started_at,
        "completed_at": completed_at,
        "finish_reason": "completed (write)",
        "error": None,
        "final_output_path": str(legacy_run_dir / "final.md"),
        "context_bundle_path": str(legacy_run_dir / "context_bundle.json"),
        "run_log_path": str(run_log_path),
    }
    # Stale API-state alias should be ignored in favor of artifact folder discovery.
    (state_dir / "legacy-run.json").write_text(
        json.dumps({"run_id": "legacy-run", **shared_payload}, ensure_ascii=True, indent=2),
        encoding="utf-8",
    )
    (state_dir / "legacy-run_01.json").write_text(
        json.dumps(
            {"run_id": "legacy-run_01", **shared_payload}, ensure_ascii=True, indent=2
        ),
        encoding="utf-8",
    )

    app = create_app(runs_dir=runs_dir, max_workers=1, markdown_dir=markdown_dir)
    with TestClient(app) as client:
        listed_runs = client.get("/api/v1/runs")
        assert listed_runs.status_code == 200
        payload = listed_runs.json()
        assert payload["total"] == 1
        assert payload["runs"][0]["run_id"] == "legacy-run_01"


def test_api_run_supports_header_api_key_override(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    runs_dir = tmp_path / "output"
    markdown_dir = tmp_path / "documents"
    markdown_dir.mkdir(parents=True, exist_ok=True)
    captured_api_key: dict[str, str | None] = {"value": None}

    def _stub_load_config(_path: Path | None = None) -> AppConfig:
        return _build_config(runs_dir=runs_dir, markdown_dir=markdown_dir)

    def _stub_run_pipeline(
        question: str,
        config: AppConfig,
        run_id: str | None = None,
        log_llm_payload: bool = True,
        analysis_mode: str = "aggregate",
        api_key_override: str | None = None,
        selected_cities: list[str] | None = None,
    ) -> RunPaths:
        assert run_id is not None
        assert analysis_mode == "aggregate"
        assert selected_cities is None
        captured_api_key["value"] = api_key_override
        return _write_success_artifacts(question=question, run_id=run_id, config=config)

    monkeypatch.setattr("backend.api.services.run_executor.load_config", _stub_load_config)
    monkeypatch.setattr("backend.api.services.run_executor.run_pipeline", _stub_run_pipeline)

    app = create_app(runs_dir=runs_dir, max_workers=1, markdown_dir=markdown_dir)
    with TestClient(app) as client:
        start = client.post(
            "/api/v1/runs",
            json={"question": "Header key test", "run_id": "run-header-key"},
            headers={"X-OpenRouter-Api-Key": "sk-or-v1-user-test-key"},
        )
        assert start.status_code == 202
        terminal = _poll_until_terminal(client, "run-header-key")
        assert terminal["status"] == "completed"

    assert captured_api_key["value"] == "sk-or-v1-user-test-key"


def test_api_key_error_is_reported_with_sanitized_message(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    runs_dir = tmp_path / "output"
    markdown_dir = tmp_path / "documents"
    markdown_dir.mkdir(parents=True, exist_ok=True)

    def _stub_load_config(_path: Path | None = None) -> AppConfig:
        return _build_config(runs_dir=runs_dir, markdown_dir=markdown_dir)

    def _failing_run_pipeline(
        question: str,
        config: AppConfig,
        run_id: str | None = None,
        log_llm_payload: bool = True,
        analysis_mode: str = "aggregate",
        api_key_override: str | None = None,
        selected_cities: list[str] | None = None,
    ) -> RunPaths:
        assert analysis_mode == "aggregate"
        assert api_key_override is None
        assert selected_cities is None
        raise RuntimeError(
            "Incorrect API key provided: sk-or-v1-abcdefghijklmnopqrstuv0123456789"
        )

    monkeypatch.setattr("backend.api.services.run_executor.load_config", _stub_load_config)
    monkeypatch.setattr(
        "backend.api.services.run_executor.run_pipeline", _failing_run_pipeline
    )

    app = create_app(runs_dir=runs_dir, max_workers=1, markdown_dir=markdown_dir)
    with TestClient(app) as client:
        start = client.post(
            "/api/v1/runs",
            json={"question": "Key fail test", "run_id": "run-key-fail"},
        )
        assert start.status_code == 202
        terminal = _poll_until_terminal(client, "run-key-fail")
        assert terminal["status"] == "failed"
        assert terminal["error"]["code"] == "API_KEY_ERROR"
        assert "sk-or-v1-" not in terminal["error"]["message"]
