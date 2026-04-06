from io import BytesIO

from docx import Document

from backend.api.services.document_export import markdown_to_docx_bytes


def test_markdown_to_docx_bytes_preserves_headings_lists_and_tables() -> None:
    markdown = """# Report

Summary with **bold** context and `inline_code`.

- First point
- Second point

| City | Status |
| --- | --- |
| Munich | On track |
| Leipzig | In review |
"""

    payload = markdown_to_docx_bytes(markdown)
    document = Document(BytesIO(payload))

    assert document.paragraphs[0].text == "Report"
    assert any(paragraph.text == "First point" for paragraph in document.paragraphs)
    assert any(paragraph.text == "Second point" for paragraph in document.paragraphs)

    assert len(document.tables) == 1
    table = document.tables[0]
    assert table.rows[0].cells[0].text == "City"
    assert table.rows[0].cells[1].text == "Status"
    assert table.rows[1].cells[0].text == "Munich"
    assert table.rows[2].cells[1].text == "In review"


def test_markdown_to_docx_bytes_excludes_inline_ref_tokens() -> None:
    markdown = (
        "# Report\n\n"
        "Summary with evidence [ref_1][ref_2].\n\n"
        "| City | Status |\n"
        "| --- | --- |\n"
        "| Munich | On track [ref_3] |\n\n"
        "Reference link [ref_guide](https://example.com/docs).\n"
    )

    payload = markdown_to_docx_bytes(markdown)
    document = Document(BytesIO(payload))

    assert document.paragraphs[1].text == "Summary with evidence."
    assert document.tables[0].rows[1].cells[1].text == "On track"
    assert document.paragraphs[2].text == "Reference link ref_guide (https://example.com/docs)."
